import docx
import sys

def read_docx(path):
    doc = docx.Document(path)
    for i, p in enumerate(doc.paragraphs):
        print(f"P{i}: {p.text}")
    for i, t in enumerate(doc.tables):
        print(f"Table {i}:")
        for j, row in enumerate(t.rows):
            print(f"  Row {j}: {[cell.text for cell in row.cells]}")

if __name__ == '__main__':
    read_docx(sys.argv[1])
