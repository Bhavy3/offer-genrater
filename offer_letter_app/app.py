import streamlit as st
import datetime
import io
import json
import os
import shutil
import subprocess
import sys

st.set_page_config(page_title="Vibrant Offer Letter", layout="wide")

st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Outfit:wght@300;400;600;800&family=Inter:wght@400;500;700&display=swap');

    :root {
        --bg-deep: #050510;
        --glass-bg: rgba(13, 14, 38, 0.45);
        --glass-border: rgba(255, 255, 255, 0.08);
        --accent-cyan: #00f2fe;
        --accent-pink: #fa709a;
        --accent-purple: #c471ed;
        --text-main: #ffffff;
        --text-muted: #a0aec0;
    }

    /* Immersive Dark Website Background */
    .stApp {
        font-family: 'Inter', sans-serif;
        background-color: var(--bg-deep);
        background-image: 
            radial-gradient(ellipse at top left, rgba(250, 112, 154, 0.15), transparent 40%),
            radial-gradient(ellipse at bottom right, rgba(0, 242, 254, 0.15), transparent 40%),
            radial-gradient(circle at 50% 50%, rgba(196, 113, 237, 0.1), transparent 50%);
        background-attachment: fixed;
        color: var(--text-main);
    }

    [data-testid="stSidebar"] {
        background: rgba(5, 5, 16, 0.8) !important;
        backdrop-filter: blur(25px);
        -webkit-backdrop-filter: blur(25px);
        border-right: 1px solid var(--glass-border);
    }

    .block-container {
        padding-top: 2.5rem;
        padding-bottom: 2.5rem;
    }

    h1, h2, h3, h4, h5, h6, .hero-title {
        font-family: 'Outfit', sans-serif !important;
    }

    /* Premium Glass Cards */
    .hero-card, .glass-card {
        background: var(--glass-bg);
        backdrop-filter: blur(20px) saturate(180%);
        -webkit-backdrop-filter: blur(20px) saturate(180%);
        border: 1px solid var(--glass-border);
        border-radius: 24px;
        box-shadow: 0 10px 40px rgba(0, 0, 0, 0.5);
        padding: 2rem;
        margin-bottom: 1.5rem;
        transition: transform 0.4s cubic-bezier(0.175, 0.885, 0.32, 1.275);
    }

    .hero-card:hover, .glass-card:hover {
        transform: translateY(-5px);
        border: 1px solid rgba(0, 242, 254, 0.3);
        box-shadow: 0 15px 50px rgba(0, 242, 254, 0.15);
    }

    .hero-badge {
        display: inline-block;
        border-radius: 999px;
        padding: 0.5rem 1rem;
        background: rgba(250, 112, 154, 0.15);
        border: 1px solid rgba(250, 112, 154, 0.4);
        color: #ff9a9e;
        font-family: 'Outfit', sans-serif;
        font-size: 0.85rem;
        font-weight: 600;
        letter-spacing: 0.15em;
        text-transform: uppercase;
        margin-bottom: 1rem;
    }

    .hero-title {
        font-size: 3.2rem !important;
        font-weight: 800 !important;
        background: linear-gradient(to right, #fff, var(--accent-cyan));
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        margin-bottom: 0.8rem !important;
        line-height: 1.1 !important;
    }

    .hero-subtitle {
        color: var(--text-muted) !important;
        font-size: 1.15rem !important;
        line-height: 1.7 !important;
        max-width: 80%;
    }

    /* Metric Boxes */
    .metric-box {
        background: rgba(20, 22, 50, 0.5);
        backdrop-filter: blur(15px);
        border: 1px solid rgba(255, 255, 255, 0.05);
        border-radius: 16px;
        padding: 1.5rem;
        margin-bottom: 1rem;
        position: relative;
        overflow: hidden;
        transition: all 0.3s ease;
    }

    .metric-box::before {
        content: '';
        position: absolute;
        top: 0; left: 0; width: 4px; height: 100%;
        background: linear-gradient(to bottom, var(--accent-cyan), var(--accent-purple));
    }

    .metric-box:hover {
        transform: scale(1.02);
        background: rgba(20, 22, 50, 0.8);
        box-shadow: 0 0 30px rgba(196, 113, 237, 0.2);
    }

    .metric-label { 
        color: #8b9eb3; 
        font-size: 0.85rem; 
        text-transform: uppercase; 
        letter-spacing: 0.1em;
        font-family: 'Outfit', sans-serif;
        font-weight: 600;
    }
    .metric-value { 
        color: #ffffff; 
        font-size: 1.8rem; 
        font-weight: 800; 
        font-family: 'Outfit', sans-serif;
        text-shadow: 0 0 15px rgba(0, 242, 254, 0.5);
    }

    /* Inputs */
    .stTextInput > div > div > input,
    .stDateInput > div > div > input {
        background-color: rgba(10, 11, 26, 0.7) !important;
        border: 1px solid rgba(255, 255, 255, 0.1) !important;
        color: #fff !important;
        border-radius: 12px !important;
        padding: 0.7rem !important;
        font-family: 'Inter', sans-serif !important;
        transition: all 0.3s ease;
    }

    .stTextInput > div > div > input:focus,
    .stDateInput > div > div > input:focus {
        border-color: var(--accent-cyan) !important;
        box-shadow: 0 0 15px rgba(0, 242, 254, 0.3) !important;
    }

    /* Vibrant Buttons */
    .stButton > button {
        background: linear-gradient(135deg, var(--accent-purple) 0%, var(--accent-cyan) 100%) !important;
        border: none !important;
        color: white !important;
        border-radius: 14px !important;
        padding: 0.8rem 1.5rem !important;
        font-family: 'Outfit', sans-serif !important;
        font-weight: 700 !important;
        font-size: 1.1rem !important;
        letter-spacing: 0.05em !important;
        transition: all 0.3s ease !important;
        box-shadow: 0 10px 20px rgba(0, 242, 254, 0.3) !important;
        text-transform: uppercase;
        width: 100%;
    }

    .stButton > button:hover {
        transform: translateY(-3px) !important;
        box-shadow: 0 15px 30px rgba(0, 242, 254, 0.5) !important;
        background: linear-gradient(135deg, var(--accent-cyan) 0%, var(--accent-purple) 100%) !important;
    }

    /* Radio Buttons Container */
    div[role="radiogroup"] {
        background: rgba(10, 11, 26, 0.5);
        padding: 20px;
        border-radius: 16px;
        border: 1px solid rgba(255, 255, 255, 0.08);
    }

    /* Alerts and Dataframes */
    div[data-testid="stAlert"] > div { 
        background: rgba(10, 11, 26, 0.8);
        backdrop-filter: blur(15px);
        border-radius: 16px; 
        border: 1px solid rgba(250, 112, 154, 0.3); 
    }
    
    [data-testid="stDataFrame"] {
        border-radius: 16px;
        overflow: hidden;
        border: 1px solid rgba(255, 255, 255, 0.08);
        background: rgba(10, 11, 26, 0.5);
    }
</style>
""", unsafe_allow_html=True)

st.markdown("""
<div class="hero-card">
  <div class="hero-badge">Premium Offer Letter Studio</div>
  <div class="hero-title">Create refined internship offers in minutes</div>
  <div class="hero-subtitle">Generate polished offer letters, store candidate records, and search them instantly by name or enrollment number.</div>
</div>
""", unsafe_allow_html=True)

col_a, col_b, col_c = st.columns(3)
with col_a:
    st.markdown("<div class='metric-box'><div class='metric-label'>Template</div><div class='metric-value'>DOCX-based</div></div>", unsafe_allow_html=True)
with col_b:
    st.markdown("<div class='metric-box'><div class='metric-label'>Search</div><div class='metric-value'>By name / ENO</div></div>", unsafe_allow_html=True)
with col_c:
    st.markdown("<div class='metric-box'><div class='metric-label'>Exports</div><div class='metric-value'>PDF + DOCX</div></div>", unsafe_allow_html=True)

main_left, _ = st.columns([1.15, 0.85])
with main_left:
    st.markdown("""
    <div class='glass-card'>
      <h3 style='margin-top:0;'>Generation Workspace</h3>
      <p style='color:#bfd2ea;'>Use the sidebar to enter candidate details, upload a custom template, and generate a refined offer letter instantly.</p>
    </div>
    """, unsafe_allow_html=True)

# ===================== DEFAULT TEMPLATE =====================
DEFAULT_TEMPLATE_PATH = "ALAN BIJO VARGHESE.docx"
CONFIRMATION_TEMPLATE_PATH = "offrer - letter - Vibrant tech lab - final.docx"
COMPLETION_TEMPLATE_PATH = "Chaitya shah  comp.docx"
HISTORY_FILE = os.path.join(os.path.dirname(__file__), "output", "offer_history.json")


def load_offer_history():
    if not os.path.exists(HISTORY_FILE):
        return []
    try:
        with open(HISTORY_FILE, "r", encoding="utf-8") as handle:
            data = json.load(handle)
        return data if isinstance(data, list) else []
    except Exception:
        return []

def save_offer_history(entry):
    history = load_offer_history()
    history.insert(0, entry)
    with open(HISTORY_FILE, "w", encoding="utf-8") as handle:
        json.dump(history[:100], handle, indent=2)

def search_offer_history(query):
    term = (query or "").strip().lower()
    if not term:
        return load_offer_history()

    results = []
    for entry in load_offer_history():
        text = " ".join([
            entry.get("student_name", ""),
            entry.get("enrollment_no", ""),
            entry.get("college_name", ""),
            entry.get("subject", ""),
        ]).lower()
        if term in text:
            results.append(entry)
    return results

main_right = st.columns([1.15, 0.85])[1]
with main_right:
    recent = load_offer_history()[:4]
    st.markdown("""
    <div class='glass-card'>
      <h3 style='margin-top:0;'>Recently Saved</h3>
      <p style='color:#bfd2ea;'>Quick access to the latest generated records.</p>
    </div>
    """, unsafe_allow_html=True)
    if recent:
        for item in recent:
            st.markdown(f"<div class='metric-box'><div class='metric-label'>{item.get('student_name', 'Unknown')}</div><div class='metric-value'>{item.get('enrollment_no', '-')}</div></div>", unsafe_allow_html=True)
    else:
        st.info("Generate an offer letter to populate this section.")

if not os.path.exists(DEFAULT_TEMPLATE_PATH):
    fallback = r"C:\Users\IQ\OneDrive\Documents\padas 1\vibrant\offer genrater\ALAN BIJO VARGHESE.docx"
    if os.path.exists(fallback):
        DEFAULT_TEMPLATE_PATH = fallback

if not os.path.exists(CONFIRMATION_TEMPLATE_PATH):
    fallback_conf = r"C:\Users\IQ\OneDrive\Documents\padas 1\vibrant\offer genrater\offrer - letter - Vibrant tech lab - final.docx"
    if os.path.exists(fallback_conf):
        CONFIRMATION_TEMPLATE_PATH = fallback_conf

if not os.path.exists(COMPLETION_TEMPLATE_PATH):
    fallback_comp = r"C:\Users\IQ\OneDrive\Documents\padas 1\vibrant\offer genrater\Chaitya shah  comp.docx"
    if os.path.exists(fallback_comp):
        COMPLETION_TEMPLATE_PATH = fallback_comp

# ===================== SIDEBAR =====================
with st.sidebar:
    st.markdown("<div class='glass-card'><h3 style='margin-top:0;'>Document Settings</h3></div>", unsafe_allow_html=True)
    doc_type = st.radio("Select Document Type", ["Offer Letter", "Confirmation Letter", "Completion Certificate"])
    st.divider()

    if doc_type == "Offer Letter":
        st.markdown("<div class='glass-card'><h3 style='margin-top:0;'>Candidate Profile</h3><p style='color:#bfd2ea; margin-bottom:0;'>Fill the details below and generate a premium offer letter in one click.</p></div>", unsafe_allow_html=True)

        student_name = st.text_input("Student Full Name", value="bhavy gajjar")
        college_name = st.text_input("College / University", value="Vibrant University")
        enrollment_no = st.text_input("Enrollment Number", value="2204030100306")

        today = datetime.date.today()
        issue_date_obj = st.date_input("Issue Date", value=today)
        start_date_obj = st.date_input("Starting Date", value=today)
        end_date_obj = st.date_input("Ending Date", value=today + datetime.timedelta(days=30))
        subject = st.text_input("Domain / Subject", value="Java")

        issue_date = issue_date_obj.strftime("%d/%m/%Y")
        start_date = start_date_obj.strftime("%d/%m/%Y")
        end_date = end_date_obj.strftime("%d/%m/%Y")

        st.info("Default template is already loaded")
        uploaded_file = st.file_uploader("Upload Different Template (Optional)", type=["docx"])

    elif doc_type == "Confirmation Letter":
        st.markdown("<div class='glass-card'><h3 style='margin-top:0;'>Confirmation Details</h3><p style='color:#bfd2ea; margin-bottom:0;'>Enter details and student list to generate a confirmation letter.</p></div>", unsafe_allow_html=True)
        
        today = datetime.date.today()
        issue_date_obj = st.date_input("Issue Date", value=today)
        issue_date = issue_date_obj.strftime("%d/%m/%Y")
        
        technology = st.text_input("Technology / Domain", value="Python & Django")
        
        st.markdown("#### Students List")
        default_data = [
            {"No": 1, "Student Name": "Varsani Jenish", "Enrollment Number": "202312101801"},
            {"No": 2, "Student Name": "", "Enrollment Number": ""},
        ]
        students_data = st.data_editor(default_data, num_rows="dynamic", use_container_width=True)
        
        uploaded_file = None

    elif doc_type == "Completion Certificate":
        st.markdown("<div class='glass-card'><h3 style='margin-top:0;'>Certificate Details</h3><p style='color:#bfd2ea; margin-bottom:0;'>Fill the details to generate a completion certificate.</p></div>", unsafe_allow_html=True)
        
        student_name = st.text_input("Student Full Name", value="Chaitya shah")
        enrollment_no = st.text_input("Enrollment Number", value="230020116060")
        subject = st.text_input("Domain / Subject", value="Python With Data Analytics")
        grade = st.text_input("Grade", value="A")
        
        today = datetime.date.today()
        issue_date_obj = st.date_input("Issue Date", value=today)
        start_date_obj = st.date_input("Starting Date", value=today - datetime.timedelta(days=30))
        end_date_obj = st.date_input("Ending Date", value=today)

        issue_date = issue_date_obj.strftime("%d/%m/%Y")
        start_date = start_date_obj.strftime("%d/%m/%Y")
        end_date = end_date_obj.strftime("%d/%m/%Y")
        
        st.info("Default template is already loaded")
        uploaded_file = st.file_uploader("Upload Different Template (Optional)", type=["docx"])

    st.divider()
    st.markdown("<div class='glass-card'><h3 style='margin-top:0;'>Search Archive</h3><p style='color:#bfd2ea; margin-bottom:0;'>Find stored offers quickly by candidate name or enrollment number.</p></div>", unsafe_allow_html=True)
    search_query = st.text_input("Name or Enrollment No", placeholder="Search saved offer records")

    if search_query:
        matches = search_offer_history(search_query)
        if matches:
            st.caption(f"{len(matches)} record(s) found")
            for item in matches[:8]:
                st.markdown(f"- **{item.get('student_name', 'Unknown')}**\n  ENO: {item.get('enrollment_no', '-')}\n  College: {item.get('college_name', '-')}", unsafe_allow_html=True)
        else:
            st.info("No saved offer record matches your search.")
    else:
        recent = load_offer_history()[:5]
        if recent:
            st.caption("Recent saved records")
            for item in recent:
                st.markdown(f"- **{item.get('student_name', 'Unknown')}**  •  ENO: {item.get('enrollment_no', '-')}", unsafe_allow_html=True)
        else:
            st.caption("No saved records yet. Generate one to populate this list.")

    st.divider()
    generate_btn = st.button(f"🚀 Generate {doc_type}", type="primary", use_container_width=True)

# ===================== GENERATION =====================
if generate_btn:
    import docx
    import tempfile

    try:
        final_filename = ""
        student_name_to_save = ""
        college_name_to_save = ""
        enrollment_no_to_save = ""
        subject_to_save = ""

        if doc_type == "Offer Letter":
            if uploaded_file is not None:
                with open("temp_template.docx", "wb") as f:
                    f.write(uploaded_file.read())
                doc = docx.Document("temp_template.docx")
            else:
                doc = docx.Document(DEFAULT_TEMPLATE_PATH)

            student_name_upper = student_name.upper()
            
            student_name_to_save = student_name
            college_name_to_save = college_name
            enrollment_no_to_save = enrollment_no
            subject_to_save = subject

            p8_handled = False
            if len(doc.paragraphs) > 8:
                p8 = doc.paragraphs[8]
                has_alan = any('ALAN' in r.text for r in p8.runs)
                if has_alan and len(p8.runs) >= 10:
                    from docx.enum.text import WD_TAB_ALIGNMENT
                    from docx.shared import Inches
                    
                    p8.paragraph_format.tab_stops.add_tab_stop(Inches(7.0), alignment=WD_TAB_ALIGNMENT.RIGHT)
                    
                    p8.runs[1].text = student_name_upper
                    p8.runs[1].bold = True
                    p8.runs[2].text = ''
                    p8.runs[3].text = ''
                    p8.runs[4].text = ''
                    p8.runs[5].text = ''
                    
                    p8.runs[6].text = '\t'
                    p8.runs[7].text = ''
                    p8.runs[8].text = ''
                    
                    for r in p8.runs:
                        if '11/05/2026' in r.text:
                            r.text = r.text.replace('11/05/2026', issue_date)
                        elif '01/05/2026' in r.text:
                            r.text = r.text.replace('01/05/2026', issue_date)
                    
                    for i in range(10, len(p8.runs)):
                        p8.runs[i].text = ''
                    p8_handled = True

            for idx, p in enumerate(doc.paragraphs):
                if idx == 8 and p8_handled:
                    continue
                    
                for r in p.runs:
                    if '11/05/2026' in r.text:
                        r.text = r.text.replace('11/05/2026', start_date)
                    if '11/08/2026' in r.text:
                        r.text = r.text.replace('11/08/2026', end_date)
                    if '2241230265' in r.text:
                        r.text = r.text.replace('2241230265', enrollment_no)
                    if 'Data Science' in r.text:
                        r.text = r.text.replace('Data Science', subject)
                
                if 'lnternship' in p.text or 'Internship' in p.text:
                    for r in p.runs:
                        if r.text == '-':
                            r.text = ''
                        if 'lnternshi' in r.text:
                            r.text = r.text.replace('lnternshi', 'Internshi')

                has_alan = any('ALAN' in r.text for r in p.runs)
                has_bijo = any('BIJO' in r.text for r in p.runs)
                has_varghese = any('VARGHESE' in r.text for r in p.runs)
                
                if has_alan and has_bijo and has_varghese:
                    replaced_combined = False
                    for r in p.runs:
                        if 'ALAN BIJO VARGHESE' in r.text:
                            r.text = r.text.replace('ALAN BIJO VARGHESE', student_name_upper)
                            r.bold = True
                            replaced_combined = True
                            break
                    
                    if not replaced_combined:
                        for r in p.runs:
                            if 'ALAN' in r.text:
                                r.text = r.text.replace('ALAN', student_name_upper)
                                r.bold = True
                            if 'BIJO' in r.text:
                                r.text = r.text.replace('BIJO', '')
                            if 'VARGHESE' in r.text:
                                r.text = r.text.replace('VARGHESE', '')
                else:
                    for r in p.runs:
                        if 'ALAN BIJO VARGHESE' in r.text:
                            r.text = r.text.replace('ALAN BIJO VARGHESE', student_name_upper)
                            r.bold = True
                        elif 'ALAN BIJO' in r.text:
                            r.text = r.text.replace('ALAN BIJO', student_name_upper)
                            r.bold = True

            final_filename = f"Offer_Letter_{student_name.replace(' ', '_')}"

        elif doc_type == "Confirmation Letter":
            # Confirmation Letter Logic
            doc = docx.Document(CONFIRMATION_TEMPLATE_PATH)
            
            student_name_to_save = "Batch Confirmation"
            college_name_to_save = "Multiple Students"
            enrollment_no_to_save = "-"
            subject_to_save = technology
            start_date = issue_date
            end_date = issue_date

            for p in doc.paragraphs:
                if '11/06/2026' in p.text or 'Python & Django' in p.text:
                    new_text = p.text.replace('11/06/2026', issue_date).replace('Python & Django', technology)
                    if len(p.runs) > 0:
                        p.runs[0].text = new_text
                        for i in range(1, len(p.runs)):
                            p.runs[i].text = ""
            
            if len(doc.tables) > 0:
                table = doc.tables[0]
                
                # Keep only the first row (header)
                for i in range(len(table.rows) - 1, 0, -1):
                    tbl = table._tbl
                    tr = table.rows[i]._tr
                    tbl.remove(tr)
                
                # Filter valid students
                valid_students = []
                for row in students_data:
                    name = str(row.get('Student Name', '')).strip()
                    if name:
                        valid_students.append(row)
                
                # Populate rows
                for idx, row in enumerate(valid_students):
                    new_row = table.add_row()
                    new_row.cells[0].text = str(idx + 1)
                    new_row.cells[1].text = str(row.get('Student Name', ''))
                    new_row.cells[2].text = str(row.get('Enrollment Number', ''))

            final_filename = f"Confirmation_Letter_{issue_date.replace('/','_')}"

        elif doc_type == "Completion Certificate":
            if uploaded_file is not None:
                with open("temp_template.docx", "wb") as f:
                    f.write(uploaded_file.read())
                doc = docx.Document("temp_template.docx")
            else:
                doc = docx.Document(COMPLETION_TEMPLATE_PATH)

            student_name_to_save = student_name
            college_name_to_save = "-"
            enrollment_no_to_save = enrollment_no
            subject_to_save = subject

            for p in doc.paragraphs:
                if 'Issued on:' in p.text and 'Enroll on:' in p.text:
                    new_text = p.text.replace('15/06/2026', issue_date).replace('230020116060', enrollment_no)
                    if len(p.runs) > 0:
                        p.runs[0].text = new_text
                        for i in range(1, len(p.runs)):
                            p.runs[i].text = ""
                elif 'This is to certify that' in p.text:
                    p.clear()
                    p.add_run("This is to certify that ")
                    r_name = p.add_run(student_name)
                    r_name.bold = True
                    p.add_run(f" successfully completed internship with grade {grade} for ")
                    r_subj = p.add_run(subject)
                    r_subj.bold = True
                    p.add_run(f" conducted by Vibrant Technology from {start_date} to {end_date} at 801, Silicon Tower, Law Garden, opp. Axis Bank, Ellisbridge, Ahmedabad,.")

            final_filename = f"Completion_Certificate_{student_name.replace(' ', '_')}"


        # Save to temp docx
        temp_dir = os.path.join(os.path.dirname(__file__), "output", "temp")
        os.makedirs(temp_dir, exist_ok=True)
        docx_path = os.path.join(temp_dir, "document.docx")
        pdf_path = os.path.join(temp_dir, "document.pdf")
        doc.save(docx_path)

        with open(docx_path, "rb") as f:
            final_docx = f.read()

        final_pdf = None
        pdf_error = None

        def convert_docx_to_pdf(source_path, destination_path, output_dir):
            if sys.platform == "win32":
                try:
                    from docx2pdf import convert
                    try:
                        import pythoncom
                        pythoncom.CoInitialize()
                    except Exception:
                        pass
                    convert(source_path, destination_path)
                    return True
                except Exception as first_error:
                    candidates = [
                        shutil.which('libreoffice'),
                        shutil.which('soffice'),
                    ]
                    for path in [r'C:\Program Files\LibreOffice\program\soffice.exe', r'C:\Program Files (x86)\LibreOffice\program\soffice.exe']:
                        if os.path.exists(path):
                            candidates.append(path)
                            
                    converter = next((path for path in candidates if path), None)
                    if not converter:
                        raise RuntimeError(f"Windows PDF conversion failed (ensure docx2pdf is installed or MS Word/LibreOffice is available). Inner Error: {first_error}") from first_error

                    subprocess.run(
                        [converter, '--headless', '--convert-to', 'pdf', source_path, '--outdir', output_dir],
                        check=True,
                        capture_output=True,
                        text=True,
                    )
                    return True
            else:
                converter = shutil.which('libreoffice') or shutil.which('soffice')
                if not converter:
                    raise FileNotFoundError('LibreOffice/soffice was not found for PDF conversion.')

                subprocess.run(
                    [converter, '--headless', '--convert-to', 'pdf', source_path, '--outdir', output_dir],
                    check=True,
                    capture_output=True,
                    text=True,
                )
                return True

        try:
            convert_docx_to_pdf(docx_path, pdf_path, temp_dir)
            with open(pdf_path, "rb") as f:
                final_pdf = f.read()
        except Exception as e:
            pdf_error = f"PDF Conversion Failed: {e}"

        # ===================== DISPLAY RESULT =====================
        st.success(f"✅ {doc_type} Generated Successfully!")

        def log_download():
            save_offer_history({
                "student_name": student_name_to_save,
                "college_name": college_name_to_save,
                "enrollment_no": enrollment_no_to_save,
                "issue_date": issue_date,
                "start_date": start_date,
                "end_date": end_date,
                "subject": subject_to_save,
                "saved_at": datetime.datetime.now().isoformat(timespec="seconds"),
            })

        if final_pdf:
            col1, col2 = st.columns([1, 1])
            with col1:
                st.download_button(
                    label="📥 Download PDF",
                    data=final_pdf,
                    file_name=f"{final_filename}.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                    on_click=log_download
                )
            with col2:
                st.download_button(
                    label="📝 Download Word (DOCX)",
                    data=final_docx,
                    file_name=f"{final_filename}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    on_click=log_download
                )
            st.info("Live PDF Preview")
            import base64
            base64_pdf = base64.b64encode(final_pdf).decode('utf-8')
            pdf_display = f'<iframe src="data:application/pdf;base64,{base64_pdf}" width="100%" height="800" type="application/pdf"></iframe>'
            st.markdown(pdf_display, unsafe_allow_html=True)
        else:
            st.download_button(
                label="📝 Download Word (DOCX)",
                data=final_docx,
                file_name=f"{final_filename}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                on_click=log_download
            )
            st.warning("⚠️ Live PDF preview failed.")
            if pdf_error:
                st.code(pdf_error)
            st.info("You can still download the perfect DOCX file using the button above.")

    except Exception as e:
        st.error(f"Error: {e}")
        st.info("Make sure the default template exists.")

else:
    st.info("Fill details on the left sidebar and click **Generate**")
    st.caption("Made for clean text replacement using native DOCX formatting")