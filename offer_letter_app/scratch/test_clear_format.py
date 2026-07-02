import docx
from docx.shared import Pt
import os
import shutil

# Copy template to bypass locks
template_path = os.path.abspath(os.path.join("..", "_Vibrant tech lab - kishan , hetansh.docx"))
temp_template = "scratch/temp_template.docx"
shutil.copy2(template_path, temp_template)

# Load document
doc = docx.Document(temp_template)

# Simulation inputs
students_data = [
    {"No": 1, "Student Name": "bhavy", "Enrollment Number": "306"},
    {"No": 2, "Student Name": "krishna", "Enrollment Number": "2415"},
    {"No": 3, "Student Name": "dev", "Enrollment Number": "204"}
]

# Find the Names textboxes
txbx_elements = doc.element.xpath('//*[local-name()="txbxContent"]')

names_txbxs = []
for txbx in txbx_elements:
    p_elements = txbx.xpath('*[local-name()="p"]')
    if p_elements:
        from docx.text.paragraph import Paragraph
        first_p = Paragraph(p_elements[0], doc._body)
        if "Student Name" in first_p.text:
            names_txbxs.append(p_elements)

s1 = students_data[0].get('Student Name', '').strip()
s2 = students_data[1].get('Student Name', '').strip()
s3 = students_data[2].get('Student Name', '').strip()

for p_elements in names_txbxs:
    p0 = Paragraph(p_elements[0], doc._body)
    p1 = Paragraph(p_elements[1], doc._body)
    
    # Reset paragraph formatting to defaults (removes 1.33" indent and 2.8x line spacing)
    p0.paragraph_format.line_spacing = None
    p0.paragraph_format.first_line_indent = None
    p0.paragraph_format.left_indent = None
    p0.paragraph_format.space_before = None
    p0.paragraph_format.space_after = None
    
    p1.paragraph_format.line_spacing = None
    p1.paragraph_format.first_line_indent = None
    p1.paragraph_format.left_indent = None
    p1.paragraph_format.space_before = None
    p1.paragraph_format.space_after = None
    
    # 1. Update Paragraph 0 (Header + Student 1 + Student 2)
    p0.clear()
    
    # Run 1: Header "Student Name"
    r_header = p0.add_run("Student Name")
    r_header.font.name = "Arial Black"
    r_header.font.size = Pt(10)
    r_header.bold = True
    
    # Spacing before Student 1 (Row 1 is aligned with Line 3)
    p0.add_run("\n\n")
    
    # Run 2: Student 1
    r_s1 = p0.add_run(s1)
    r_s1.font.name = "Arial"
    r_s1.font.size = Pt(12)
    
    # Spacing before Student 2 (Row 2 is aligned with Line 5)
    p0.add_run("\n\n")
    
    # Run 3: Student 2
    r_s2 = p0.add_run(s2)
    r_s2.font.name = "Arial"
    r_s2.font.size = Pt(12)
    
    # 2. Update Paragraph 1 (Student 3)
    p1.clear()
    
    # Spacing before Student 3 (Row 3 is aligned with Line 7)
    p1.add_run("\n")
    
    # Run 4: Student 3
    r_s3 = p1.add_run(s3)
    r_s3.font.name = "Arial"
    r_s3.font.size = Pt(12)

# Save
doc.save("scratch/generated_clear.docx")
print("Saved rebuilt document!")

# Convert to PDF
print("Converting to PDF...")
try:
    from docx2pdf import convert
    convert("scratch/generated_clear.docx", "scratch/generated_clear.pdf")
    print("PDF converted successfully!")
except Exception as e:
    print("docx2pdf failed:", e)

# Render to PNG
print("Rendering PDF to PNG...")
import fitz
doc_pdf = fitz.open("scratch/generated_clear.pdf")
page = doc_pdf.load_page(0)
pix = page.get_pixmap(dpi=150)
pix.save("scratch/test_clear.png")
print("Saved page to scratch/test_clear.png")
